# /// script
# requires-python = ">=3.11"
# dependencies = ["python-docx>=1.2.0"]
# ///

"""Extract CubeRoot algorithm-table data from a curated DOCX.

The extractor is intentionally strict: a visually reviewed document is still the
source of truth, while this script turns its table structure into deterministic JSON.
"""

from __future__ import annotations

import argparse
import hashlib
import html
import json
import re
import sys
from dataclasses import dataclass
from pathlib import Path, PurePosixPath
from typing import Any, Iterable

from docx import Document
from docx.oxml.ns import qn
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run


GREEN_VALUES = {"00B050"}
METRICS_RE = re.compile(
    r"^\s*([^,]+?)\s*,\s*(\d+)\s*,\s*(\d+)\s*,\s*(\d+)\s*,\s*(\d+)\s*,\s*(\d+)\s*$"
)
COUNT_NOTE_RE = re.compile(r"\s*\(\d+[A-Za-z]?\*\)\s*$")
KNOWN_NOTE_START_RE = re.compile(r"^\s*\((odd)\)\s*", re.IGNORECASE)
KNOWN_NOTE_END_RE = re.compile(
    r"\s*\((put side light on F|if no swap CP)\)\s*$", re.IGNORECASE
)
ARROWS = {"->": "right", "→": "right", "<-": "left", "←": "left"}

# Content hashes make the mapping independent of Word's imageNN filenames.
ICON_TAG_BY_SHA256 = {
    "15c2e01d21985c335acbbc692e13d02374ef01a610b096768e1bb08d9b61d896": "oh",
    "94bb7a0fa163e485057b0d3bc1169beb6e397526536583161528c5766ed536fb": "ft",
    "0527997f9a6208e63a5f3f47d0323b4a5bed3f25d10a5c721e13eca2c5351fa1": "big",
    "3683a86a322b3115bfbdcf471c4ec0d1ce4d67a2a48b998240b1cc45587ef4a1": "key",
    "26c5c4dfeaa39c188f03acd9107ea65bb7c0702c3035a56ccd0d0d50855a5fc1": "fmc",
}


class ExtractError(ValueError):
    pass


@dataclass(frozen=True)
class StyledText:
    text: str
    bold: bool
    italic: bool
    underline: bool
    strike: bool
    superscript: bool
    subscript: bool


def normalized_text(value: str) -> str:
    return " ".join(value.replace("\xa0", " ").split())


def direct_run_color(run: Run) -> str | None:
    rpr = run._r.find(qn("w:rPr"))
    if rpr is not None:
        color = rpr.find(qn("w:color"))
        if color is not None:
            value = color.get(qn("w:val"))
            if value:
                return value.upper()
    rgb = run.font.color.rgb
    return str(rgb).upper() if rgb is not None else None


def is_green(run: Run) -> bool:
    return direct_run_color(run) in GREEN_VALUES


def media_targets(run: Run) -> Iterable[tuple[str, bytes]]:
    seen: set[str] = set()
    for element in run._r.iter():
        if element.tag.rsplit("}", 1)[-1] not in {"blip", "imagedata"}:
            continue
        rel_id = next(
            (
                value
                for key, value in element.attrib.items()
                if key.rsplit("}", 1)[-1] in {"embed", "id"}
            ),
            None,
        )
        if not rel_id or rel_id in seen or rel_id not in run.part.rels:
            continue
        seen.add(rel_id)
        rel = run.part.rels[rel_id]
        target = PurePosixPath(rel.target_ref).name
        try:
            blob = rel.target_part.blob
        except AttributeError:
            continue
        yield target, blob


def run_style(run: Run, text: str) -> StyledText:
    return StyledText(
        text=text,
        bold=bool(run.bold),
        italic=bool(run.italic),
        underline=bool(run.underline),
        strike=bool(run.font.strike),
        superscript=bool(run.font.superscript),
        subscript=bool(run.font.subscript),
    )


def slice_segments(segments: list[StyledText], start: int, end: int) -> list[StyledText]:
    result: list[StyledText] = []
    cursor = 0
    for segment in segments:
        next_cursor = cursor + len(segment.text)
        left = max(start, cursor)
        right = min(end, next_cursor)
        if left < right:
            result.append(
                StyledText(
                    text=segment.text[left - cursor : right - cursor],
                    bold=segment.bold,
                    italic=segment.italic,
                    underline=segment.underline,
                    strike=segment.strike,
                    superscript=segment.superscript,
                    subscript=segment.subscript,
                )
            )
        cursor = next_cursor
    return result


def render_segment(segment: StyledText) -> str:
    value = html.escape(segment.text)
    wrappers = (
        (segment.bold, "strong"),
        (segment.italic, "em"),
        (segment.underline, "u"),
        (segment.strike, "s"),
        (segment.superscript, "sup"),
        (segment.subscript, "sub"),
    )
    for enabled, tag in wrappers:
        if enabled:
            value = f"<{tag}>{value}</{tag}>"
    return value


def formula_bounds(raw: str) -> tuple[int, int, list[str]]:
    start = len(raw) - len(raw.lstrip())
    end = len(raw.rstrip())
    notes: list[str] = []

    if raw[start:end].startswith("="):
        start += 1
        while start < end and raw[start].isspace():
            start += 1

    prefix = KNOWN_NOTE_START_RE.match(raw[start:end])
    if prefix:
        notes.append(prefix.group(1))
        start += prefix.end()

    while True:
        visible = raw[start:end]
        count_suffix = COUNT_NOTE_RE.search(visible)
        if count_suffix:
            notes.insert(0, count_suffix.group(0).strip()[1:-1])
            end = start + count_suffix.start()
            continue
        prose_suffix = KNOWN_NOTE_END_RE.search(visible)
        if prose_suffix:
            notes.append(prose_suffix.group(1))
            end = start + prose_suffix.start()
            continue
        break

    while end > start and raw[end - 1].isspace():
        end -= 1
    return start, end, notes


def extract_formula(paragraph: Paragraph) -> tuple[dict[str, Any] | None, str | None, list[str]]:
    raw_all = "".join(run.text for run in paragraph.runs)
    arrow = ARROWS.get(normalized_text(raw_all))
    unknown_icons: list[str] = []
    tags: list[str] = []
    segments: list[StyledText] = []

    for run in paragraph.runs:
        for target, blob in media_targets(run):
            digest = hashlib.sha256(blob).hexdigest()
            tag = ICON_TAG_BY_SHA256.get(digest)
            if tag and tag not in tags:
                tags.append(tag)
            elif not tag:
                unknown_icons.append(f"{target}:{digest}")
        if run.text and not is_green(run):
            segments.append(run_style(run, run.text.replace("\xa0", " ")))

    raw_visible = "".join(segment.text for segment in segments)
    start, end, notes = formula_bounds(raw_visible)
    if start >= end:
        return None, arrow, unknown_icons
    formula_segments = slice_segments(segments, start, end)
    alg = "".join(segment.text for segment in formula_segments)
    alg_html = "".join(render_segment(segment) for segment in formula_segments)
    escaped_alg = html.escape(alg)

    result: dict[str, Any] = {"alg": alg}
    if alg_html != escaped_alg:
        result["algHtml"] = alg_html
    if tags:
        result["tags"] = tags
    if notes:
        result["sourceNotes"] = notes
    return result, arrow, unknown_icons


def metrics_in_cell(cell: _Cell) -> tuple[str, dict[str, int]] | None:
    for paragraph in cell.paragraphs:
        match = METRICS_RE.match(normalized_text(paragraph.text))
        if not match:
            continue
        name, etm, optimal_etm, htm, stm, atm = match.groups()
        return name, {
            "etm": int(etm),
            "optimalEtm": int(optimal_etm),
            "optimalHtm": int(htm),
            "optimalStm": int(stm),
            "optimalAtm": int(atm),
        }
    return None


def case_number(cell: _Cell) -> int:
    for paragraph in cell.paragraphs:
        text = normalized_text(paragraph.text)
        match = re.fullmatch(r"0*(\d+)", text)
        if match:
            return int(match.group(1))
    raise ExtractError(f"missing case number in cell {cell.text!r}")


def table_score(table: Table) -> int:
    return sum(
        1
        for row in table.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
        if METRICS_RE.match(normalized_text(paragraph.text))
    )


def select_table(document: Any, requested: int | None) -> tuple[int, Table]:
    if requested is not None:
        try:
            return requested, document.tables[requested]
        except IndexError as exc:
            raise ExtractError(f"table index {requested} is out of range") from exc
    scored = [(table_score(table), index, table) for index, table in enumerate(document.tables)]
    score, index, table = max(scored, default=(0, -1, None), key=lambda item: item[0])
    if table is None or score == 0:
        raise ExtractError("no table contains a five-number metric row")
    return index, table


def heading_for_row(row: Any) -> str | None:
    candidates: list[str] = []
    seen_cells: set[int] = set()
    for cell in row.cells:
        key = id(cell._tc)
        if key in seen_cells:
            continue
        seen_cells.add(key)
        value = next(
            (
                normalized_text(paragraph.text)
                for paragraph in cell.paragraphs
                if normalized_text(paragraph.text)
            ),
            "",
        )
        if value and not re.fullmatch(r"0*\d+", value):
            candidates.append(value)
    return candidates[0] if len(set(candidates)) == 1 else None


def extract(path: Path, table_index: int | None, expected_cases: int | None) -> dict[str, Any]:
    document = Document(path)
    selected_index, table = select_table(document, table_index)
    category: str | None = None
    categories: list[str] = []
    cases: list[dict[str, Any]] = []
    unknown_icons: set[str] = set()

    for row_index, row in enumerate(table.rows):
        metric_sides = [metrics_in_cell(row.cells[column]) for column in (0, 2)]
        if not any(metric_sides):
            heading = heading_for_row(row)
            if heading:
                category = heading
                if heading not in categories:
                    categories.append(heading)
            continue
        if not category:
            raise ExtractError(f"row {row_index} contains a case before any category heading")

        row_cases: dict[str, dict[str, Any]] = {}
        row_arrows: dict[str, set[str]] = {"left": set(), "right": set()}
        for side, info_column, formula_column in (("left", 0, 1), ("right", 2, 3)):
            metric_data = metrics_in_cell(row.cells[info_column])
            if not metric_data:
                continue
            name, metrics = metric_data
            no = case_number(row.cells[info_column])
            formulas: list[dict[str, Any]] = []
            for paragraph in row.cells[formula_column].paragraphs:
                formula, arrow, paragraph_unknown = extract_formula(paragraph)
                unknown_icons.update(paragraph_unknown)
                if arrow:
                    row_arrows[side].add(arrow)
                if formula:
                    formulas.append(formula)
            if not formulas:
                raise ExtractError(f"case {no} has no non-green formulas")
            item = {
                "no": no,
                "name": name,
                "category": category,
                "position": len(cases),
                "metrics": metrics,
                "algs": formulas,
            }
            cases.append(item)
            row_cases[side] = item

        for side, arrows in row_arrows.items():
            if not arrows:
                continue
            expected = "right" if side == "left" else "left"
            if arrows != {expected} or set(row_cases) != {"left", "right"}:
                raise ExtractError(f"unsupported arrow relation at row {row_index}: {side}={sorted(arrows)}")
            other = "right" if side == "left" else "left"
            row_cases[side]["scrambleFrom"] = row_cases[other]["no"]

    numbers = [item["no"] for item in cases]
    duplicates = sorted({number for number in numbers if numbers.count(number) > 1})
    if duplicates:
        raise ExtractError(f"duplicate case numbers: {duplicates}")
    if expected_cases is not None and len(cases) != expected_cases:
        raise ExtractError(f"expected {expected_cases} cases, extracted {len(cases)}")
    if unknown_icons:
        raise ExtractError("unknown formula icons: " + ", ".join(sorted(unknown_icons)))

    return {
        "source": str(path),
        "tableIndex": selected_index,
        "categories": categories,
        "caseCount": len(cases),
        "cases": cases,
    }


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("docx", type=Path)
    parser.add_argument("--table", type=int, default=None, help="zero-based table index")
    parser.add_argument("--expect-cases", type=int, default=None)
    parser.add_argument("--output", type=Path, default=None)
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    try:
        result = extract(args.docx.resolve(), args.table, args.expect_cases)
    except (ExtractError, OSError) as exc:
        print(f"error: {exc}", file=sys.stderr)
        return 2
    payload = json.dumps(result, ensure_ascii=False, indent=2) + "\n"
    if args.output:
        args.output.parent.mkdir(parents=True, exist_ok=True)
        args.output.write_text(payload, encoding="utf-8", newline="\n")
    else:
        sys.stdout.write(payload)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
